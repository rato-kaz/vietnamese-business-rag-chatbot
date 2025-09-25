import os
import docx
from typing import Dict, List, Any, Optional
import re
from pathlib import Path


class TemplateParser:
    def __init__(self, templates_dir: str = "templates"):
        """Initialize template parser for business registration forms."""
        self.templates_dir = Path(templates_dir)
        self.form_fields = {}
        self.templates = {}
        
        # Load and parse all template files
        self._load_templates()
    
    def _load_templates(self):
        """Load all template files and extract form fields."""
        if not self.templates_dir.exists():
            print(f"Templates directory not found: {self.templates_dir}")
            return
        
        template_files = [
            "danh_sach_chu_so_huu.docx",
            "danh_sach_co_dong.docx", 
            "dieu_le_cong_ty.docx",
            "giay_de_nghi.docx",
            "giay_uy_quyen.docx"
        ]
        
        for template_file in template_files:
            file_path = self.templates_dir / template_file
            if file_path.exists():
                try:
                    content = self._load_docx_content(file_path)
                    fields = self._extract_form_fields(content, template_file)
                    self.templates[template_file] = {
                        "content": content,
                        "fields": fields
                    }
                    print(f"Loaded template: {template_file} with {len(fields)} fields")
                except Exception as e:
                    print(f"Error loading template {template_file}: {e}")
    
    def _load_docx_content(self, file_path: Path) -> str:
        """Load content from docx file."""
        try:
            doc = docx.Document(file_path)
            content = ""
            for paragraph in doc.paragraphs:
                content += paragraph.text + "\n"
            
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        content += cell.text + "\n"
            
            return content
        except Exception as e:
            print(f"Error reading docx file {file_path}: {e}")
            return ""
    
    def _extract_form_fields(self, content: str, template_name: str) -> List[Dict[str, Any]]:
        """Extract form fields from template content."""
        fields = []
        
        # Common patterns for form fields in Vietnamese documents
        patterns = [
            # Fields with brackets or underlines
            r'([^:\n]+?):\s*[_\.\s]{3,}',  # Field: _______
            r'([^:\n]+?):\s*\[[\s_]*\]',   # Field: [     ]
            r'([^:\n]+?):\s*\.{3,}',       # Field: ......
            
            # Fields in specific templates
            r'Tên công ty:\s*[_\.\s]*',
            r'Địa chỉ:\s*[_\.\s]*',
            r'Người đại diện:\s*[_\.\s]*',
            r'Vốn điều lệ:\s*[_\.\s]*',
            r'Ngành nghề kinh doanh:\s*[_\.\s]*',
            
            # Owner/shareholder information
            r'Họ và tên:\s*[_\.\s]*',
            r'CMND/CCCD:\s*[_\.\s]*',
            r'Ngày sinh:\s*[_\.\s]*',
            r'Số điện thoại:\s*[_\.\s]*',
            r'Email:\s*[_\.\s]*',
            r'Tỷ lệ góp vốn:\s*[_\.\s]*',
        ]
        
        # Extract based on template type
        if "danh_sach_chu_so_huu" in template_name:
            fields.extend(self._extract_owner_fields(content))
        elif "danh_sach_co_dong" in template_name:
            fields.extend(self._extract_shareholder_fields(content))
        elif "dieu_le_cong_ty" in template_name:
            fields.extend(self._extract_company_charter_fields(content))
        elif "giay_de_nghi" in template_name:
            fields.extend(self._extract_application_fields(content))
        elif "giay_uy_quyen" in template_name:
            fields.extend(self._extract_authorization_fields(content))
        
        return fields
    
    def _extract_owner_fields(self, content: str) -> List[Dict[str, Any]]:
        """Extract fields for owner list."""
        return [
            {
                "field_name": "chu_so_huu_ho_ten",
                "display_name": "Họ và tên chủ sở hữu",
                "field_type": "text",
                "required": True,
                "description": "Họ và tên đầy đủ của chủ sở hữu công ty"
            },
            {
                "field_name": "chu_so_huu_cccd",
                "display_name": "Số CMND/CCCD",
                "field_type": "text",
                "required": True,
                "description": "Số chứng minh nhân dân hoặc căn cước công dân"
            },
            {
                "field_name": "chu_so_huu_ngay_sinh",
                "display_name": "Ngày sinh",
                "field_type": "date",
                "required": True,
                "description": "Ngày sinh của chủ sở hữu (dd/mm/yyyy)"
            },
            {
                "field_name": "chu_so_huu_dia_chi",
                "display_name": "Địa chỉ thường trú",
                "field_type": "text",
                "required": True,
                "description": "Địa chỉ thường trú của chủ sở hữu"
            },
            {
                "field_name": "chu_so_huu_dien_thoai",
                "display_name": "Số điện thoại",
                "field_type": "text",
                "required": False,
                "description": "Số điện thoại liên hệ"
            }
        ]
    
    def _extract_shareholder_fields(self, content: str) -> List[Dict[str, Any]]:
        """Extract fields for shareholder list."""
        return [
            {
                "field_name": "co_dong_ho_ten",
                "display_name": "Họ và tên cổ đông",
                "field_type": "text",
                "required": True,
                "description": "Họ và tên đầy đủ của cổ đông"
            },
            {
                "field_name": "co_dong_cccd",
                "display_name": "Số CMND/CCCD",
                "field_type": "text",
                "required": True,
                "description": "Số chứng minh nhân dân hoặc căn cước công dân"
            },
            {
                "field_name": "co_dong_ty_le_von",
                "display_name": "Tỷ lệ góp vốn (%)",
                "field_type": "number",
                "required": True,
                "description": "Tỷ lệ phần trăm góp vốn của cổ đông"
            },
            {
                "field_name": "co_dong_gia_tri_von",
                "display_name": "Giá trị vốn góp",
                "field_type": "number",
                "required": True,
                "description": "Giá trị vốn góp bằng tiền (VNĐ)"
            }
        ]
    
    def _extract_company_charter_fields(self, content: str) -> List[Dict[str, Any]]:
        """Extract fields for company charter."""
        return [
            {
                "field_name": "ten_cong_ty",
                "display_name": "Tên công ty",
                "field_type": "text",
                "required": True,
                "description": "Tên đầy đủ của công ty"
            },
            {
                "field_name": "ten_cong_ty_tieng_anh",
                "display_name": "Tên công ty (tiếng Anh)",
                "field_type": "text",
                "required": False,
                "description": "Tên công ty bằng tiếng Anh (nếu có)"
            },
            {
                "field_name": "dia_chi_tru_so",
                "display_name": "Địa chỉ trụ sở chính",
                "field_type": "text",
                "required": True,
                "description": "Địa chỉ đầy đủ của trụ sở chính công ty"
            },
            {
                "field_name": "von_dieu_le",
                "display_name": "Vốn điều lệ",
                "field_type": "number",
                "required": True,
                "description": "Vốn điều lệ của công ty (VNĐ)"
            },
            {
                "field_name": "nganh_nghe_kinh_doanh",
                "display_name": "Ngành nghề kinh doanh",
                "field_type": "text",
                "required": True,
                "description": "Mô tả chi tiết các ngành nghề kinh doanh"
            }
        ]
    
    def _extract_application_fields(self, content: str) -> List[Dict[str, Any]]:
        """Extract fields for application form."""
        return [
            {
                "field_name": "nguoi_dai_dien",
                "display_name": "Người đại diện pháp luật",
                "field_type": "text",
                "required": True,
                "description": "Họ tên người đại diện pháp luật"
            },
            {
                "field_name": "chuc_vu_dai_dien",
                "display_name": "Chức vụ",
                "field_type": "text",
                "required": True,
                "description": "Chức vụ của người đại diện (Giám đốc, Tổng Giám đốc, ...)"
            },
            {
                "field_name": "ngay_lap_don",
                "display_name": "Ngày lập đơn",
                "field_type": "date",
                "required": True,
                "description": "Ngày lập đơn đăng ký (dd/mm/yyyy)"
            }
        ]
    
    def _extract_authorization_fields(self, content: str) -> List[Dict[str, Any]]:
        """Extract fields for authorization letter."""
        return [
            {
                "field_name": "nguoi_uy_quyen",
                "display_name": "Người ủy quyền",
                "field_type": "text",
                "required": True,
                "description": "Họ tên người ủy quyền"
            },
            {
                "field_name": "nguoi_duoc_uy_quyen",
                "display_name": "Người được ủy quyền",
                "field_type": "text",
                "required": True,
                "description": "Họ tên người được ủy quyền"
            },
            {
                "field_name": "noi_dung_uy_quyen",
                "display_name": "Nội dung ủy quyền",
                "field_type": "text",
                "required": True,
                "description": "Mô tả cụ thể những việc được ủy quyền"
            }
        ]
    
    def get_all_form_fields(self) -> Dict[str, List[Dict[str, Any]]]:
        """Get all form fields from all templates."""
        all_fields = {}
        for template_name, template_data in self.templates.items():
            all_fields[template_name] = template_data["fields"]
        return all_fields
    
    def get_template_fields(self, template_name: str) -> List[Dict[str, Any]]:
        """Get fields for a specific template."""
        if template_name in self.templates:
            return self.templates[template_name]["fields"]
        return []
    
    def get_required_fields(self) -> List[Dict[str, Any]]:
        """Get all required fields across all templates."""
        required_fields = []
        for template_data in self.templates.values():
            for field in template_data["fields"]:
                if field.get("required", False):
                    required_fields.append(field)
        return required_fields
    
    def generate_form_collection_questions(self) -> List[Dict[str, Any]]:
        """Generate a list of questions to collect form data from user."""
        questions = []
        
        # Get unique required fields
        seen_fields = set()
        for template_data in self.templates.values():
            for field in template_data["fields"]:
                field_name = field["field_name"]
                if field_name not in seen_fields and field.get("required", False):
                    seen_fields.add(field_name)
                    
                    question = {
                        "field_name": field_name,
                        "question": f"Vui lòng nhập {field['display_name'].lower()}:",
                        "field_type": field["field_type"],
                        "description": field.get("description", ""),
                        "required": field.get("required", False)
                    }
                    questions.append(question)
        
        return questions
    
    def validate_field_value(self, field_name: str, value: str) -> tuple[bool, str]:
        """Validate a field value based on its type."""
        # Find field definition
        field_def = None
        for template_data in self.templates.values():
            for field in template_data["fields"]:
                if field["field_name"] == field_name:
                    field_def = field
                    break
            if field_def:
                break
        
        if not field_def:
            return True, ""  # Unknown field, assume valid
        
        field_type = field_def.get("field_type", "text")
        
        if field_type == "date":
            # Validate date format dd/mm/yyyy
            if not re.match(r'^\d{2}/\d{2}/\d{4}$', value):
                return False, "Định dạng ngày không đúng. Vui lòng nhập theo format dd/mm/yyyy"
        
        elif field_type == "number":
            try:
                float(value.replace(",", "").replace(".", ""))
            except ValueError:
                return False, "Giá trị phải là số"
        
        elif field_type == "text":
            if field_def.get("required", False) and not value.strip():
                return False, "Trường này là bắt buộc"
        
        return True, ""