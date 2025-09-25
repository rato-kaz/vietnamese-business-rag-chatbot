import os
import re
import yaml
from typing import List, Dict, Any, Optional
from datetime import datetime
import docx
from docx.document import Document as DocxDocument
from langchain.schema import Document


class DocumentProcessor:
    def __init__(self, config_path: str = "config/config.yaml"):
        """Initialize the document processor with configuration."""
        with open(config_path, 'r', encoding='utf-8') as file:
            self.config = yaml.safe_load(file)
        
        self.chunk_size = self.config['document_processing']['chunk_size']
        self.chunk_overlap = self.config['document_processing']['chunk_overlap']
        self.supported_formats = self.config['document_processing']['supported_formats']
        
        # Document type mappings
        self.document_type_mapping = {
            'luật': 'Luật',
            'nghị định': 'Nghị định',
            'thông tư': 'Thông tư',
            'quyết định': 'Quyết định',
            'tt': 'Thông tư',
            'nđ-cp': 'Nghị định',
            'qđ': 'Quyết định'
        }
        
        # Agency mappings
        self.agency_mapping = {
            'qh': 'Quốc hội',
            'ttg': 'Thủ tướng Chính phủ',
            'btc': 'Bộ Tài chính',
            'bkhđt': 'Bộ Kế hoạch và Đầu tư',
            'cp': 'Chính phủ'
        }
    
    def extract_metadata_from_filename(self, filename: str) -> Dict[str, Any]:
        """Extract metadata from filename pattern."""
        metadata = {
            "source": filename,
            "source_file": os.path.splitext(filename)[0],
            "document_number": None,
            "document_type": None,
            "issue_date": None,
            "issuing_agency": None,
            "effective_date": None,
            "expiry_date": None,
            "confidential_level": "Công khai",
            "issue_year": None,
            "law_field": "khac"
        }
        
        # Extract from patterns like "_130_2017_TT-BTC_373081.docx"
        pattern1 = r"_(\d+)_(\d{4})_([A-Z\-]+)_\d+\.docx"
        match1 = re.search(pattern1, filename)
        if match1:
            number, year, agency_code = match1.groups()
            metadata["document_number"] = f"{number}/{year}/{agency_code}"
            metadata["issue_year"] = int(year)
            
            # Determine document type and agency from code
            agency_lower = agency_code.lower()
            for key, value in self.agency_mapping.items():
                if key in agency_lower:
                    metadata["issuing_agency"] = value
                    break
                    
            if 'tt' in agency_lower:
                metadata["document_type"] = "Thông tư"
            elif 'nd' in agency_lower:
                metadata["document_type"] = "Nghị định"
            elif 'qd' in agency_lower:
                metadata["document_type"] = "Quyết định"
        
        # Extract from patterns like "Luật-03-2022-QH15.docx"
        pattern2 = r"(Luật|Nghị định|Thông tư|Quyết định)-(\d+)-(\d{4})-([A-Z0-9]+)\.docx"
        match2 = re.search(pattern2, filename)
        if match2:
            doc_type, number, year, agency_code = match2.groups()
            metadata["document_type"] = doc_type
            metadata["document_number"] = f"{number}/{year}/{agency_code}"
            metadata["issue_year"] = int(year)
            
            agency_lower = agency_code.lower()
            for key, value in self.agency_mapping.items():
                if key in agency_lower:
                    metadata["issuing_agency"] = value
                    break
        
        return metadata
    
    def extract_metadata_from_content(self, content: str, base_metadata: Dict[str, Any]) -> Dict[str, Any]:
        """Extract additional metadata from document content."""
        metadata = base_metadata.copy()
        
        # Extract document title (usually in first few lines)
        lines = content.split('\n')[:10]
        for line in lines:
            line = line.strip()
            if len(line) > 20 and not re.match(r'^(Điều|Chương|\d+\.)', line):
                if not metadata.get("document_title"):
                    metadata["document_title"] = line
                    break
        
        # Extract dates
        date_patterns = [
            r"ngày\s+(\d{1,2})\s+tháng\s+(\d{1,2})\s+năm\s+(\d{4})",
            r"(\d{1,2})/(\d{1,2})/(\d{4})"
        ]
        
        for pattern in date_patterns:
            matches = re.findall(pattern, content, re.IGNORECASE)
            if matches:
                day, month, year = matches[0]
                try:
                    date_str = f"{day:0>2}/{month:0>2}/{year}"
                    if not metadata.get("issue_date"):
                        metadata["issue_date"] = date_str
                except:
                    pass
        
        return metadata
    
    def chunk_by_articles(self, content: str, metadata: Dict[str, Any]) -> List[Document]:
        """Chunk document by articles (Điều) and sections."""
        chunks = []
        
        # Split by articles (Điều)
        article_pattern = r'(Điều\s+\d+[.\s]*[^\n]*)'
        articles = re.split(article_pattern, content)
        
        current_article = None
        current_article_title = None
        
        for i, section in enumerate(articles):
            if re.match(r'Điều\s+\d+', section.strip()):
                current_article = section.strip()
                current_article_title = section.strip()
            elif section.strip() and current_article:
                # Further split by numbered items (1., 2., a., b., etc.)
                item_pattern = r'(\n\s*[0-9a-z]+\.\s*)'
                items = re.split(item_pattern, section)
                
                current_item_content = ""
                current_khoan_code = None
                
                for j, item in enumerate(items):
                    if re.match(r'\n\s*[0-9a-z]+\.\s*', item):
                        # Save previous item if exists
                        if current_item_content.strip():
                            chunk_metadata = metadata.copy()
                            chunk_metadata.update({
                                "article_code": current_article,
                                "dieu_code": current_article,
                                "dieu_title": current_article_title,
                                "chunk_title": f"{current_article}" + (f" - {current_khoan_code}" if current_khoan_code else ""),
                                "khoan_code": current_khoan_code,
                                "entity_type": "article_section"
                            })
                            
                            chunks.append(Document(
                                page_content=current_item_content.strip(),
                                metadata=chunk_metadata
                            ))
                        
                        # Start new item
                        current_khoan_code = item.strip()
                        current_item_content = ""
                    else:
                        current_item_content += item
                
                # Add last item
                if current_item_content.strip():
                    chunk_metadata = metadata.copy()
                    chunk_metadata.update({
                        "article_code": current_article,
                        "dieu_code": current_article,
                        "dieu_title": current_article_title,
                        "chunk_title": f"{current_article}" + (f" - {current_khoan_code}" if current_khoan_code else ""),
                        "khoan_code": current_khoan_code,
                        "entity_type": "article_section"
                    })
                    
                    chunks.append(Document(
                        page_content=current_item_content.strip(),
                        metadata=chunk_metadata
                    ))
        
        return chunks
    
    def load_docx_document(self, file_path: str) -> str:
        """Load content from a .docx file."""
        try:
            doc = docx.Document(file_path)
            content = ""
            for paragraph in doc.paragraphs:
                content += paragraph.text + "\n"
            return content
        except Exception as e:
            print(f"Error loading {file_path}: {str(e)}")
            return ""
    
    def process_directory(self, directory_path: str) -> List[Document]:
        """Process all supported documents in a directory."""
        all_documents = []
        
        for root, dirs, files in os.walk(directory_path):
            for file in files:
                if file.endswith('.docx'):
                    file_path = os.path.join(root, file)
                    try:
                        # Load document content
                        content = self.load_docx_document(file_path)
                        if not content.strip():
                            continue
                        
                        # Extract metadata
                        base_metadata = self.extract_metadata_from_filename(file)
                        metadata = self.extract_metadata_from_content(content, base_metadata)
                        
                        # Chunk by articles
                        chunks = self.chunk_by_articles(content, metadata)
                        
                        if not chunks:  # Fallback to simple chunking
                            chunk_metadata = metadata.copy()
                            chunk_metadata.update({
                                "article_code": None,
                                "dieu_code": None,
                                "dieu_title": None,
                                "chunk_title": metadata.get("document_title", file),
                                "khoan_code": None,
                                "entity_type": "document"
                            })
                            
                            chunks = [Document(
                                page_content=content,
                                metadata=chunk_metadata
                            )]
                        
                        all_documents.extend(chunks)
                        print(f"Processed: {file} -> {len(chunks)} chunks")
                        
                    except Exception as e:
                        print(f"Error processing {file_path}: {str(e)}")
        
        return all_documents
    
    def get_document_stats(self, documents: List[Document]) -> Dict[str, Any]:
        """Get statistics about processed documents."""
        stats = {
            'total_documents': len(documents),
            'total_characters': sum(len(doc.page_content) for doc in documents),
            'average_chunk_size': sum(len(doc.page_content) for doc in documents) / len(documents) if documents else 0,
            'document_types': {},
            'agencies': {},
            'years': {}
        }
        
        for doc in documents:
            metadata = doc.metadata
            
            # Count document types
            doc_type = metadata.get('document_type')
            if doc_type:
                stats['document_types'][doc_type] = stats['document_types'].get(doc_type, 0) + 1
            
            # Count agencies
            agency = metadata.get('issuing_agency')
            if agency:
                stats['agencies'][agency] = stats['agencies'].get(agency, 0) + 1
            
            # Count years
            year = metadata.get('issue_year')
            if year:
                stats['years'][str(year)] = stats['years'].get(str(year), 0) + 1
        
        return stats